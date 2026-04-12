# -*- coding: utf-8 -*-
"""只補五、六章的更新（三四章已在上一個腳本完成）"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

path = r'D:\APP\Nice to meet you\設計文件\遊戲內程式設計.docx'
doc = Document(path)

style_map = {}
for s in doc.styles:
    if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name not in style_map:
        style_map[s.name] = s

# ── 五、年齡池測試說明更新 ─────────────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if '年齡池測試' in p.text and p.style.name.startswith('Heading'):
        desc = doc.paragraphs[i + 1]
        desc.clear()
        desc.add_run(
            'Debug Panel 看天數：Day 3 起切少年池（card_010~012），Day 6 起切青年池（card_002~003），依此類推。'
            '（測試閾值，正式版切換天數不同）'
        )
        print('五章 年齡池測試說明 已更新')
        break

# ── 六、上線前清單更新 ───────────────────────────────────────
for p in doc.paragraphs:
    if '補充足夠數量的各年齡池牌卡' in p.text:
        p.clear()
        p.add_run('補充足夠數量的各年齡池牌卡（每池至少涵蓋對應天數 × 3 張）')
        print('六章 牌卡數量清單項 已更新')
        break

for i, p in enumerate(doc.paragraphs):
    if '實作晨間殘留系統' in p.text:
        new_item = doc.add_paragraph(
            'AGE_POOL_THRESHOLDS 改回正式版數值（juvenile:10, youth:30, midlife:60, elder:90）',
            style=style_map.get('List Bullet')
        )
        p._element.addnext(new_item._element)
        print('六章 新增閾值還原清單項')
        break

doc.save(path)
print('Done')
