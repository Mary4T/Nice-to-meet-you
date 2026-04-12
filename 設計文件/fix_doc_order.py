# -*- coding: utf-8 -*-
"""清除倒序插入的三四章內容，重新正確插入"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

path = r'D:\APP\Nice to meet you\設計文件\遊戲內程式設計.docx'
doc = Document(path)

style_map = {}
for s in doc.styles:
    if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name not in style_map:
        style_map[s.name] = s

body = doc.element.body

# ── 找四章和五章的 element，三章插入內容夾在它們之間 ────────────
three_elem = None
four_elem  = None
five_elem  = None
for p in doc.paragraphs:
    if '三、已實作功能' in p.text: three_elem = p._element
    if '四、重要常數'   in p.text: four_elem  = p._element
    if '五、測試流程'   in p.text: five_elem  = p._element

# 刪除三到四之間（不含三和四標題本身）的所有段落
deleting = False
to_del = []
for child in list(body):
    if child is three_elem:
        deleting = True
        continue
    if child is four_elem:
        deleting = False
        break
    if deleting:
        to_del.append(child)
for elem in to_del:
    body.remove(elem)
print(f'三章舊內容已刪除（{len(to_del)} 個元素）')

# 刪除四到五之間的所有段落
deleting = False
to_del2 = []
for child in list(body):
    if child is four_elem:
        deleting = True
        continue
    if child is five_elem:
        deleting = False
        break
    if deleting:
        to_del2.append(child)
for elem in to_del2:
    body.remove(elem)
print(f'四章舊內容已刪除（{len(to_del2)} 個元素）')

# ── 重新插入（用 addprevious + 正序，插在四章標題之前）─────────
def insert_before(anchor_elem, text, style_name):
    p = doc.add_paragraph(text, style=style_map.get(style_name))
    anchor_elem.addprevious(p._element)

rows_3 = [
    ('', 'Normal'),
    ('年齡池切換系統', 'Heading 2'),
    ('【測試階段設定】', 'Heading 3'),
    ('五個年齡池（童年／少年／青年／中年／老年），閾值壓縮供開發測試用：', 'Normal'),
    ('童年 → 第 0 天起', 'Normal'),
    ('少年 → 第 3 天起（正式版：第 10 天）', 'Normal'),
    ('青年 → 第 6 天起（正式版：第 30 天）', 'Normal'),
    ('中年 → 第 9 天起（正式版：第 60 天）', 'Normal'),
    ('老年 → 第 12 天起（正式版：第 90 天）', 'Normal'),
    ('【正式版設定】', 'Heading 3'),
    ('各池需備有足夠牌數（每池至少涵蓋該期間天數 × 3 張）。', 'Normal'),
    ('【切換機制說明】', 'Heading 3'),
    ('換日時（advanceDay）自動比對新天數與閾值，若跨越閾值則將新年齡池的所有牌洗牌後附加至 mainQueue 末端。', 'Normal'),
    ('舊年齡池尚未打完的牌不移除，繼續留在 queue 中直到被抽到。', 'Normal'),
    ('每個年齡池的牌只在切換時加入一次，不重複補充。', 'Normal'),
    ('【目前各池牌數（測試階段）】', 'Heading 3'),
    ('童年：5 張（card_001, 006~009）', 'Normal'),
    ('少年：3 張（card_010~012）', 'Normal'),
    ('青年：2 張（card_002~003）', 'Normal'),
    ('中年：2 張（card_004~005）', 'Normal'),
    ('老年：0 張（待補充）', 'Normal'),
    ('', 'Normal'),
]

for text, style in rows_3:
    insert_before(four_elem, text, style)
print('三章 正確插入完成')

# 四章內容插在五章標題之前
rows_4 = [
    ('', 'Normal'),
    ('AGE_POOL_THRESHOLDS（年齡池切換閾值）', 'Heading 3'),
    ('定義在 src/core/constants/gameConfig.ts。', 'Normal'),
    ('測試階段：{ childhood:0, juvenile:3, youth:6, midlife:9, elder:12 }', 'Normal'),
    ('正式版本：{ childhood:0, juvenile:10, youth:30, midlife:60, elder:90 }', 'Normal'),
    ('上線前須將閾值改回正式版數值，並確認各池牌數足夠。', 'Normal'),
    ('', 'Normal'),
]

for text, style in rows_4:
    insert_before(five_elem, text, style)
print('四章 正確插入完成')

doc.save(path)
print('Done')
