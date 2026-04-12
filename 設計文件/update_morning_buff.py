# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

path = r'D:\APP\Nice to meet you\設計文件\補充設計決策.docx'
doc = Document(path)

# 建立 style 名稱 → style 物件的對照表（避免 duplicate name 問題）
style_map = {}
for s in doc.styles:
    if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name not in style_map:
        style_map[s.name] = s

def find_heading_index(doc, keyword):
    for i, p in enumerate(doc.paragraphs):
        if keyword in p.text and p.style and p.style.name.startswith('Heading'):
            return i
    return -1

def set_table_borders(table):
    """套用與文件原有表格相同的格式：固定寬度 9026dxa、auto 色邊框、tblLook 全 0"""
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 移除舊的 tblW / tblBorders / tblLook（如果有）
    for tag in ('w:tblW', 'w:tblBorders', 'w:tblLook', 'w:tblCellMar'):
        el = tblPr.find(qn(tag))
        if el is not None:
            tblPr.remove(el)

    # tblW
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9026')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # tblBorders
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # tblCellMar
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ('left', 'right'):
        mar = OxmlElement(f'w:{side}')
        mar.set(qn('w:w'), '10')
        mar.set(qn('w:type'), 'dxa')
        tblCellMar.append(mar)
    tblPr.append(tblCellMar)

    # tblLook
    tblLook = OxmlElement('w:tblLook')
    tblLook.set(qn('w:val'), '0000')
    tblLook.set(qn('w:firstRow'), '0')
    tblLook.set(qn('w:lastRow'), '0')
    tblLook.set(qn('w:firstColumn'), '0')
    tblLook.set(qn('w:lastColumn'), '0')
    tblLook.set(qn('w:noHBand'), '0')
    tblLook.set(qn('w:noVBand'), '0')
    tblPr.append(tblLook)

# 找「三、晨間殘留」
start = find_heading_index(doc, '三、晨間殘留')
# 找下一個同級 Heading（四、...）
end = find_heading_index(doc, '四、牌堆三層架構')

if start == -1:
    print('找不到晨間殘留段落')
    exit(1)

# 刪除 start+1 到 end-1 的段落（保留標題本身）
to_delete = doc.paragraphs[start+1:end]
for p in to_delete:
    p._element.getparent().remove(p._element)

# 重新找位置（刪除後索引改變）
start = find_heading_index(doc, '三、晨間殘留')

# 新增更新後的內容（倒序插入，因為每次都插在 start 後面）
new_content = [
    ('3-1 機制概述', 'Heading 3'),
    ('夢境卡結束後，玩家的選擇結果會以晨間殘留的形式影響隔天的裏數值效果。晨間殘留持續一天，天結束後（advanceDay）自動清除。', None),
    ('3-2 分數計算', 'Heading 3'),
    ('每張夢境卡的每個方向選項各帶一個 morningScore 欄位（範圍 -3 到 +3）。', None),
    ('三張夢境卡打完後，三個 morningScore 加總，總分範圍為 -9 到 +9。', None),
    ('3-3 分數區間與效果等級', 'Heading 3'),
    ('', None),  # placeholder for table
    ('3-4 乘數作用範圍', 'Heading 3'),
    ('晨間殘留乘數只作用於以下五個裏數值：Stability（穩定）、Drive（積極）、Logic（理性）、Emotion（感性）、hide Vitality（裏生命值）。', None),
    ('正面效果（值為正數）乘以「正面乘數」；負面效果（值為負數）乘以「負面乘數」。結果四捨五入為整數。', None),
    ('Entropy（壓抑熵）、Hypocrisy（偽善度）、Ruthlessness（冷血度）、聲望、金錢不受晨間殘留影響。', None),
    ('3-5 壓力牌', 'Heading 3'),
    ('Debuff 狀態下，當天有機率從【獨立壓力牌牌庫】額外插入一張難度較高的牌。壓力牌牌庫待劇本設計師設計，目前實作先留 TODO。', None),
]

# 倒序插入（這樣順序才正確）
ref_idx = start
for text, style_name in reversed(new_content):
    ref = doc.paragraphs[ref_idx]._element
    style_obj = style_map.get(style_name) if style_name else None
    p = doc.add_paragraph(text, style=style_obj)
    ref.addnext(p._element)

# 現在插入表格（在 3-3 標題後面的 placeholder）
idx_33 = find_heading_index(doc, '3-3 分數區間')
placeholder = doc.paragraphs[idx_33 + 1]

tbl_data = [
    ['總分', '等級', '正面效果乘數', '負面效果乘數', '壓力牌'],
    ['-9 ~ -7', '重度 Debuff', '× 0.7', '× 1.3', '高機率'],
    ['-6 ~ -4', '中度 Debuff', '× 0.8', '× 1.2', '中機率'],
    ['-3 ~ -1', '輕度 Debuff', '× 0.9', '× 1.1', '低機率'],
    ['0',        '無效果',     '× 1.0', '× 1.0', '無'],
    ['1 ~ 3',   '輕度 Buff',   '× 1.1', '× 0.9', '無'],
    ['4 ~ 6',   '中度 Buff',   '× 1.2', '× 0.8', '無'],
    ['7 ~ 9',   '重度 Buff',   '× 1.3', '× 0.7', '無'],
]

table = doc.add_table(rows=len(tbl_data), cols=5)
for i, row_data in enumerate(tbl_data):
    for j, cell_text in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = cell_text
        if i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

set_table_borders(table)

# 把表格移到 placeholder 後面，刪除 placeholder
placeholder._element.addnext(table._element)
placeholder._element.getparent().remove(placeholder._element)

doc.save(path)
print('Done')
