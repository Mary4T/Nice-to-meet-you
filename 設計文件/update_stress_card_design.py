# -*- coding: utf-8 -*-
"""
更新補充設計決策.docx：
1. 將 3-4 的「四捨五入」改為「無條件進位」
2. 展開 3-5 壓力牌系統的完整設計（取代佔位文字）
3. 修復段落 67 中夾帶的「四、牌堆三層架構」標題文字，還原為獨立 Heading 1
4. 從六、待討論事項移除已確定的壓力牌機率項目
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from lxml import etree

path = r'D:\APP\Nice to meet you\設計文件\補充設計決策.docx'
doc = Document(path)

# ── 建立 style_map ──────────────────────────────────────────
style_map = {}
for s in doc.styles:
    if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name not in style_map:
        style_map[s.name] = s

body = doc.element.body

# ── 1. 3-4 四捨五入 → 無條件進位 ────────────────────────────
for p in doc.paragraphs:
    if '四捨五入為整數' in p.text:
        for run in p.runs:
            if '四捨五入為整數' in run.text:
                run.text = run.text.replace('四捨五入為整數', '無條件進位為整數')
        print('3-4 已更新')
        break

# ── 2. 找 3-5 壓力牌 heading 與其後的段落 ───────────────────
# 找到 3-5 heading index
stress_heading_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') and '3-5' in p.text and '壓力牌' in p.text:
        stress_heading_idx = i
        break

print(f'3-5 heading at paragraph index {stress_heading_idx}')

# 找到緊接在 3-5 後面的舊佔位段落（paragraphs[stress_heading_idx + 1]）
old_placeholder = doc.paragraphs[stress_heading_idx + 1]
print(f'舊內容: {old_placeholder.text[:60]}')

# 這個段落夾帶了「四、牌堆三層架構」，我們要：
# a) 只保留 3-5 的新內容段落
# b) 在所有新段落後插入獨立的 Heading 1 for 四

# ── 3. 替換舊佔位段落，插入新內容 ───────────────────────────
ref_elem = old_placeholder._element

new_paras_content = [
    # (text, style_name)
    ('壓力牌在換日時（新的一天開始前）判定是否插入當天牌序。壓力牌插入而非替換原有牌卡，當天總牌數增加。格式與一般牌卡相同（四向選項）。', 'Normal'),
    ('3-5-1 觸發機率與上限', 'Heading 4'),
    ('輕度 Debuff：30% 機率，當天最多 1 張。', 'Normal'),
    ('中度 Debuff：60% 機率，當天最多 1 張。', 'Normal'),
    ('重度 Debuff：第一張 80% 機率；觸發後第二張再獨立判定 80% 機率，當天最多 2 張。', 'Normal'),
    ('3-5-2 出現位置', 'Heading 4'),
    ('每張壓力牌在今天三張正常牌之間的位置獨立抽籤：slot 1（60%）、slot 2（25%）、slot 3（15%）。', 'Normal'),
    ('重度 Debuff 兩張壓力牌的位置：第一張正常抽，第二張從剩餘兩個 slot 重新按比例分配。', 'Normal'),
    ('3-5-3 牌組分級', 'Heading 4'),
    ('三個 Debuff 等級各自對應獨立壓力牌組（輕/中/重），選牌時從對應等級隨機抽取。', 'Normal'),
    ('現有內容：輕度 3 張、中度 4 張、重度 3 張，共 10 張初稿，後續由劇本設計師擴充。', 'Normal'),
]

# 先刪除舊段落
body.remove(ref_elem)

# 找新的插入位置：現在 3-5 heading 後面就是 4-1 heading（或四的 Heading 1）
# 我們要在 3-5 heading 之後、四章之前插入
# 重新取得 3-5 heading element
stress_heading_elem = doc.paragraphs[stress_heading_idx]._element

# 倒序插入，讓每個 addnext 都排在正確位置
for text, style_name in reversed(new_paras_content):
    style_obj = style_map.get(style_name)
    p = doc.add_paragraph(text, style=style_obj)
    stress_heading_elem.addnext(p._element)

# 插入「四、牌堆三層架構」獨立 Heading 1（之前被塞在舊段落裡）
h1 = doc.add_paragraph('四、牌堆三層架構（Deck Engine）', style=style_map.get('Heading 1'))
# 放在所有新段落的最後（倒序中最先放的是第一個）
# 需要放在新插入的最後一段之後，也就是 3-5 所有新段落之後
# 先找 4-1 heading
para_41 = None
for p in doc.paragraphs:
    if '4-1' in p.text and p.style.name.startswith('Heading'):
        para_41 = p
        break
if para_41:
    para_41._element.addprevious(h1._element)
    print('四、牌堆三層架構 Heading 1 已插入')

# ── 4. 從六、待討論事項移除已確定的壓力牌機率項目 ────────────
for p in doc.paragraphs:
    if '壓力牌的具體觸發機率數值' in p.text:
        body.remove(p._element)
        print('六、待討論 — 壓力牌機率項目已移除')
        break

doc.save(path)
print('Done')
