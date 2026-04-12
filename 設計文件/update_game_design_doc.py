# -*- coding: utf-8 -*-
"""更新遊戲內程式設計.docx：補充年齡池切換設計、測試 vs 完整版對照"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

path = r'D:\APP\Nice to meet you\設計文件\遊戲內程式設計.docx'
doc = Document(path)

style_map = {}
for s in doc.styles:
    if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name not in style_map:
        style_map[s.name] = s

body = doc.element.body

def add_after(ref_para, text, style_name):
    style_obj = style_map.get(style_name)
    p = doc.add_paragraph(text, style=style_obj)
    ref_para._element.addnext(p._element)
    return p

# ── 1. 三、已實作功能 vs 完整版對照（段落 13）────────────────
# 在第 14 段（空行）後插入內容
ref = doc.paragraphs[14]  # 三章後的空行

rows = [
    ('年齡池切換系統', 'Heading 2'),
    ('', 'Normal'),
    ('【測試階段設定】', 'Heading 3'),
    ('五個年齡池（童年／少年／青年／中年／老年），切換閾值壓縮供開發測試：', 'Normal'),
    ('童年 → 第 0 天起', 'Normal'),
    ('少年 → 第 3 天起（正式版：第 10 天）', 'Normal'),
    ('青年 → 第 6 天起（正式版：第 30 天）', 'Normal'),
    ('中年 → 第 9 天起（正式版：第 60 天）', 'Normal'),
    ('老年 → 第 12 天起（正式版：第 90 天）', 'Normal'),
    ('', 'Normal'),
    ('【正式版設定】', 'Heading 3'),
    ('切換閾值依設計文件補充設計決策 1-3 節定義，各池需備有足夠牌數（每池至少涵蓋該期間天數 × 3 張）。', 'Normal'),
    ('', 'Normal'),
    ('【切換機制說明】', 'Heading 3'),
    ('換日時（advanceDay）自動比對新天數與閾值，若跨越閾值則將新年齡池的所有牌洗牌後附加至 mainQueue 末端。', 'Normal'),
    ('舊年齡池尚未打完的牌不移除，繼續留在 queue 中直到被抽到。', 'Normal'),
    ('每個年齡池的牌只在切換時加入一次，不重複補充。', 'Normal'),
    ('', 'Normal'),
    ('【目前各池牌數（測試階段）】', 'Heading 3'),
    ('童年：5 張（card_001, 006~009）', 'Normal'),
    ('少年：3 張（card_010~012）', 'Normal'),
    ('青年：2 張（card_002~003）', 'Normal'),
    ('中年：2 張（card_004~005）', 'Normal'),
    ('老年：0 張（待補充）', 'Normal'),
    ('', 'Normal'),
]

# 倒序插入，讓 addnext 保持正確順序
for text, style_name in reversed(rows):
    add_after(ref, text, style_name)

print('三章 年齡池切換 已插入')

# ── 2. 四、重要常數（段落 15 往後，插入點需重新取得）───────────
# 重新取段落（因為插入後 index 已改變）
four_idx = None
for i, p in enumerate(doc.paragraphs):
    if '四、重要常數' in p.text:
        four_idx = i
        break

ref4 = doc.paragraphs[four_idx + 1]  # 四章後的空行

constant_rows = [
    ('AGE_POOL_THRESHOLDS（年齡池切換閾值）', 'Heading 3'),
    ('定義在 src/core/constants/gameConfig.ts。', 'Normal'),
    ('測試階段：{ childhood:0, juvenile:3, youth:6, midlife:9, elder:12 }', 'Normal'),
    ('正式版本：{ childhood:0, juvenile:10, youth:30, midlife:60, elder:90 }', 'Normal'),
    ('上線前須將閾值改回正式版數值，並確認各池牌數足夠。', 'Normal'),
    ('', 'Normal'),
]

for text, style_name in reversed(constant_rows):
    add_after(ref4, text, style_name)

print('四章 重要常數 已插入')

# ── 3. 五-年齡池測試（更新說明）───────────────────────────────
for i, p in enumerate(doc.paragraphs):
    if '年齡池測試' in p.text and p.style.name.startswith('Heading'):
        # 找其後的說明段落
        old_desc = doc.paragraphs[i + 1]
        old_desc.text = (
            'Debug Panel 看天數：Day 3 起切少年池（card_010~012），Day 6 起切青年池（card_002~003），依此類推。'
            '（測試閾值，正式版切換天數不同）'
        )
        print('五章 年齡池測試說明 已更新')
        break

# ── 4. 六、上線前必做清單（更新舊項目、新增項目）──────────────
for p in doc.paragraphs:
    if '補充足夠數量的各年齡池牌卡' in p.text:
        p.text = '補充足夠數量的各年齡池牌卡（每池至少涵蓋對應天數 × 3 張）'
        print('六章 牌卡數量清單項 已更新')
        break

for p in doc.paragraphs:
    if '實作晨間殘留系統' in p.text:
        # 在此項後插入新項目
        new_item = doc.add_paragraph(
            'AGE_POOL_THRESHOLDS 改回正式版數值（juvenile:10, youth:30, midlife:60, elder:90）',
            style=style_map.get('List Bullet')
        )
        p._element.addnext(new_item._element)
        print('六章 新增閾值還原清單項')
        break

doc.save(path)
print('Done')
