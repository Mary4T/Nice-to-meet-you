# -*- coding: utf-8 -*-
"""
更新補充設計決策.docx：
1. 新增第十一章（十一、故事鏈系統）11-0 ~ 11-6
2. 在 7-6 後續事件欄補充填寫說明
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

PATH = r'D:\APP\Nice to meet you\設計文件\補充設計決策.docx'
doc = Document(PATH)

# ── 輔助函式 ──────────────────────────────────────────────────
def get_body():
    return doc.element.body

def find_para_by_text(keyword):
    """找到第一個包含 keyword 的段落元素"""
    for p in doc.element.body.findall('.//' + qn('w:p')):
        text = ''.join(r.text or '' for r in p.findall('.//' + qn('w:t')))
        if keyword in text:
            return p
    return None

def clone_heading_style(ref_para):
    """複製參考段落的樣式名稱"""
    pPr = ref_para.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            return pStyle.get(qn('w:val'))
    return None

def add_para_after(ref_elem, text, style=None, bold=False, keep_ref=None):
    """在 ref_elem 後插入一個新段落，回傳新段落元素"""
    new_p = OxmlElement('w:p')
    # pPr
    pPr = OxmlElement('w:pPr')
    if style:
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style)
        pPr.append(pStyle)
    new_p.append(pPr)
    # run
    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        new_p.append(r)
    ref_elem.addnext(new_p)
    return new_p

def add_table_after(ref_elem, headers, rows):
    """在 ref_elem 後插入一個簡單表格，回傳 tbl 元素"""
    from docx.oxml.table import CT_Tbl
    tbl = OxmlElement('w:tbl')
    # tblPr
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    tbl.append(tblPr)

    def make_row(cells, is_header=False):
        tr = OxmlElement('w:tr')
        for cell_text in cells:
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            if is_header:
                b = OxmlElement('w:b')
                rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = str(cell_text) if cell_text is not None else ''
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        return tr

    tbl.append(make_row(headers, is_header=True))
    for row in rows:
        tbl.append(make_row(row))

    ref_elem.addnext(tbl)
    return tbl

# ══════════════════════════════════════════════════════════════
# 1. 更新 7-6：在「後續事件」說明後補充填寫規則
# ══════════════════════════════════════════════════════════════
para_76 = find_para_by_text('7-6')
if para_76:
    # 找到 7-6 後面的表格，找到「後續事件」所在的行後面
    # 直接在 7-6 段落後找到含「後續事件」文字的段落或表格，
    # 並在整個 7-6 區段末尾加一個補充說明段落
    # 策略：找到 7-6 段落本身，然後找接下來最近的段落有「後續事件」的，
    #        在那之後插入說明
    elem = para_76
    target = None
    for _ in range(30):
        elem = elem.getnext()
        if elem is None:
            break
        text = ''.join(r.text or '' for r in elem.findall('.//' + qn('w:t')))
        if '後續事件' in text and ('空白' in text or '故事鏈' in text or '排程' in text):
            target = elem
            break
    if target is not None:
        note_p = add_para_after(
            target,
            '【填寫規則】後續事件欄填入故事鏈的起始卡ID（如 sc_001），'
            '不是鏈ID（如 chain_001）。鏈ID 僅作為 Excel 管理分組用途。',
            bold=False
        )
        print('7-6 補充說明已插入')
    else:
        print('WARNING: 找不到 7-6 後續事件行，跳過')
else:
    print('WARNING: 找不到 7-6 段落')

# ══════════════════════════════════════════════════════════════
# 2. 新增第十一章：在文件末尾（第十章之後）插入
# ══════════════════════════════════════════════════════════════
# 找到最後一個實體段落，在其後依序插入（用 addnext 逆序插入）
# 為保持順序，改為「每次找最新插入的元素當 ref」

def find_last_body_elem():
    elems = list(doc.element.body)
    # 排除 sectPr
    for e in reversed(elems):
        if e.tag != qn('w:sectPr'):
            return e
    return doc.element.body[-1]

def append_section(ref_holder, text, style=None, bold=False):
    """在 ref_holder[0]（最後插入的元素）後插入，並更新 ref_holder[0]"""
    new_elem = add_para_after(ref_holder[0], text, style=style, bold=bold)
    ref_holder[0] = new_elem
    return new_elem

def append_table(ref_holder, headers, rows):
    new_elem = add_table_after(ref_holder[0], headers, rows)
    ref_holder[0] = new_elem
    return new_elem

# 找十章最後位置
para_10 = find_para_by_text('10-6')
if para_10 is None:
    para_10 = find_para_by_text('10-5')
if para_10 is None:
    para_10 = find_para_by_text('十、條件選項')
if para_10 is None:
    para_10 = find_last_body_elem()

# 往後走，找到十章的最後一個元素（遇到 sectPr 或下一個大章節停止）
elem = para_10
while True:
    nxt = elem.getnext()
    if nxt is None or nxt.tag == qn('w:sectPr'):
        break
    text = ''.join(r.text or '' for r in nxt.findall('.//' + qn('w:t')))
    # 如果遇到新的大章節標題（十一以後）就停
    if nxt.tag == qn('w:p') and any(k in text for k in ['十一、', '十二、']):
        break
    elem = nxt

ref = [elem]  # 用 list 讓 append_section 可以修改

# ── 十一、故事鏈系統 ─────────────────────────────────────────

# 先插入一個空行分隔
append_section(ref, '')

# 大標題
append_section(ref, '十一、故事鏈系統（Story Chain System）', style='Heading1')

# 11-0
append_section(ref, '11-0 概述', style='Heading2')
append_section(ref,
    '故事鏈是從普通牌的特定選項觸發的連續分支事件序列。'
    '觸發後暫時接管牌面顯示（storyChain 緩衝區），'
    '打完最後一張鏈內牌後自動接回主牌池正常流程。'
    '故事鏈對應牌堆三層架構中的第二層（故事鏈緩衝區）。'
)

# 11-1
append_section(ref, '11-1 觸發方式', style='Heading2')
append_section(ref,
    '目前支援：選項觸發。'
    '在「卡牌資料.xlsx」的「後續事件」欄填入故事鏈起始卡ID，'
    '玩家選擇該選項後立即進入故事鏈。'
)
append_section(ref,
    '注意：後續事件欄填的是起始卡ID（如 sc_001），'
    '不是鏈ID（如 chain_001）。'
    '鏈ID 僅作 Excel 分組管理用，不用於程式觸發邏輯。'
)
append_section(ref,
    '暫不支援：數值閾值觸發（未來可擴充）。'
)

# 11-2
append_section(ref, '11-2 鏈內牌結構', style='Heading2')
append_section(ref,
    '每張故事鏈牌有四個方向（上/下/左/右），每個方向包含以下欄位：'
)
append_table(ref,
    headers=['欄位', '說明'],
    rows=[
        ['text',    '顯示給玩家的選項文字'],
        ['effects', '效果（同普通牌格式）'],
        ['next',    '下一張鏈內牌的卡ID；空白代表鏈在此結束'],
        ['endsDay', 'TRUE = 這個選項結束今日，觸發 advanceDay；FALSE = 不換日'],
    ]
)

# 11-3
append_section(ref, '11-3 分支邏輯', style='Heading2')
append_section(ref,
    '同一張鏈內牌的不同方向可以指向不同的下一張牌，形成分支樹狀結構。'
    '例如：上/下/左指向 sc_002，右指向 sc_003。'
    '玩家的選擇決定鏈的走向，不同分支可有不同情境與效果。'
)

# 11-4
append_section(ref, '11-4 天數計算', style='Heading2')
append_section(ref,
    '鏈內牌不計入主牌池的每日 3 張計數（cardsPlayedToday 不變）。'
    '換日規則由 endsDay 欄位獨立控制：'
)
append_section(ref,
    '・endsDay = TRUE：選擇後呼叫 advanceDay，同步觸發年齡池切換檢查。'
)
append_section(ref,
    '・endsDay = FALSE：不換日，直接跳到下一張鏈內牌。'
)
append_section(ref,
    '夢境卡觸發規則：Entropy 達 100 時不在鏈內立即觸發，'
    '等整條鏈結束後才插入夢境卡序列（詳見 2-5）。'
)

# 11-5
append_section(ref, '11-5 鏈結束後的行為', style='Heading2')
append_section(ref,
    '當某個方向的 next 為空時，鏈正式結束（storyChain 設為 null）。'
    '系統從主牌池 mainQueue 拉下一張普通牌並顯示，'
    '同時計入今日牌數（cardsPlayedToday + 1）。'
    '若此時 endsDay = TRUE，則在拉牌前先呼叫 startNewDay 重置當日計數。'
)

# 11-6
append_section(ref, '11-6 Excel 管理格式', style='Heading2')
append_section(ref,
    '故事鏈牌存放於「卡牌資料.xlsx」的「故事鏈牌」sheet，欄位如下：'
)
append_table(ref,
    headers=['欄位', '說明'],
    rows=[
        ['鏈ID',   '同一條鏈的所有牌填相同值（僅作分組識別，不影響程式邏輯）'],
        ['卡ID',   '此牌的唯一 ID（如 sc_001）；同一張牌的後續方向行留空'],
        ['情境',   '顯示給玩家的情境文字；同一張牌的後續方向行留空'],
        ['方向',   '上/下/左/右'],
        ['選項文字', '該方向的選項說明'],
        ['下一張', '跳往的下一張卡ID；空白 = 鏈結束'],
        ['結束今日', 'TRUE / FALSE'],
        ['效果欄', '聲望～靈魂完整度（格式同普通牌卡資料）'],
    ]
)
append_section(ref,
    '匯入腳本：import_story_chain.py，輸出至 story_chain_cards.json。'
    '普通牌觸發設定：在「卡牌資料」sheet 的「後續事件」欄填入起始卡ID，'
    '由 import_cards.py 轉換為 storyChainStart 欄位。'
)

doc.save(PATH)
print('Done — 補充設計決策.docx 已更新（新增第十一章）')
