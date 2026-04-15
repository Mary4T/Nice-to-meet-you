# -*- coding: utf-8 -*-
"""更新補充設計決策.docx 的 1-3 年齡池切換閾值表格：
   - 加入少年池
   - 更新天數（正式版 0/10/30/60/90）
   - 新增「測試值」欄（0/3/6/9/12）
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

PATH = r'D:\APP\Nice to meet you\設計文件\補充設計決策.docx'
doc = Document(PATH)

# ── 找到 1-3 的表格 ───────────────────────────────────────────
# 策略：找含有「1-3」標題的段落，再往後找第一個 Table
target_table = None
found_section = False

for i, block in enumerate(doc.element.body):
    tag = block.tag.split('}')[-1]
    if tag == 'p':
        text = ''.join(r.text for r in block.findall('.//' + qn('w:t')))
        if '1-3' in text and '年齡池' in text:
            found_section = True
    if found_section and tag == 'tbl':
        target_table = block
        break

if target_table is None:
    print('ERROR: 找不到 1-3 的表格')
    sys.exit(1)

print('找到表格，開始更新...')

# ── 取得 Table 物件 ───────────────────────────────────────────
from docx.table import Table
table = None
for t in doc.tables:
    if t._tbl is target_table:
        table = t
        break

if table is None:
    print('ERROR: 無法取得 Table 物件')
    sys.exit(1)

print(f'目前行數: {len(table.rows)}, 欄數: {len(table.columns)}')

# ── 定義新表格資料 ─────────────────────────────────────────────
NEW_DATA = [
    ('年齡池',   '正式版切換天數', '測試版切換天數', '備註'),
    ('童年 Childhood', '第 0 天',  '第 0 天',   '遊戲開始'),
    ('少年 Juvenile',  '第 10 天', '第 3 天',   '可調整'),
    ('青年 Youth',     '第 30 天', '第 6 天',   '可調整'),
    ('中年 Midlife',   '第 60 天', '第 9 天',   '可調整'),
    ('老年 Elder',     '第 90 天', '第 12 天',  '可調整'),
]

# ── 重建表格 ──────────────────────────────────────────────────
# 因為原表格是 3 欄，新的是 4 欄，最乾淨的方式是：
# 1. 清空所有行的儲存格文字
# 2. 如果行數不夠就用 XML 複製最後一行補齊
# 3. 如果欄數不夠就在每行末尾加欄

def get_cell_text(cell):
    return cell.text

def set_cell_text(cell, text, bold=False):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold

def add_column_to_row(row_elem, template_cell_elem):
    """在行尾新增一個儲存格（複製模板儲存格的格式）"""
    new_tc = copy.deepcopy(template_cell_elem)
    # 清空文字
    for t in new_tc.findall('.//' + qn('w:t')):
        t.text = ''
    row_elem.append(new_tc)
    return new_tc

def add_row(table, template_row):
    """在表格末尾新增一行（複製模板行）"""
    new_tr = copy.deepcopy(template_row)
    table._tbl.append(new_tr)
    return new_tr

# 先確保有 4 欄：如果現有行的儲存格 < 4，補一欄
current_cols = len(table.columns)
if current_cols < 4:
    for row in table.rows:
        template_cell = row.cells[-1]._tc
        add_column_to_row(row._tr, template_cell)
    print(f'已將欄數從 {current_cols} 擴充到 4')

# 確保有 6 行（1 標題 + 5 資料）
current_rows = len(table.rows)
needed_rows = len(NEW_DATA)
if current_rows < needed_rows:
    template_row = table.rows[-1]._tr
    for _ in range(needed_rows - current_rows):
        add_row(table, template_row)
    print(f'已將行數從 {current_rows} 擴充到 {needed_rows}')

# 寫入資料
for r_idx, row_data in enumerate(NEW_DATA):
    row = table.rows[r_idx]
    is_header = (r_idx == 0)
    for c_idx, cell_text in enumerate(row_data):
        if c_idx < len(row.cells):
            set_cell_text(row.cells[c_idx], cell_text, bold=is_header)

print('表格資料寫入完成')

# 如果表格行數比需要的多，刪除多餘行
while len(table.rows) > needed_rows:
    table._tbl.remove(table.rows[-1]._tr)

doc.save(PATH)
print('Done — 補充設計決策.docx 已更新')
