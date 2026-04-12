# -*- coding: utf-8 -*-
"""補充三、四章的年齡池切換內容"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

path = r'D:\APP\Nice to meet you\設計文件\遊戲內程式設計.docx'
doc = Document(path)

style_map = {}
for s in doc.styles:
    if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name not in style_map:
        style_map[s.name] = s

def insert_after(ref_elem, text, style_name):
    p = doc.add_paragraph(text, style=style_map.get(style_name))
    ref_elem.addnext(p._element)
    return p._element

# ── 三、已實作功能：找空行後插入 ─────────────────────────────
three_blank = None
for i, p in enumerate(doc.paragraphs):
    if '三、已實作功能' in p.text:
        three_blank = doc.paragraphs[i + 1]._element
        break

rows_3 = [
    ('年齡池切換系統', 'Heading 2'),
    ('【測試階段設定】', 'Heading 3'),
    ('五個年齡池（童年／少年／青年／中年／老年），閾值壓縮供開發測試用：', 'Normal'),
    ('童年 → 第 0 天起', 'Normal'),
    ('少年 → 第 3 天起（正式版：第 10 天）', 'Normal'),
    ('青年 → 第 6 天起（正式版：第 30 天）', 'Normal'),
    ('中年 → 第 9 天起（正式版：第 60 天）', 'Normal'),
    ('老年 → 第 12 天起（正式版：第 90 天）', 'Normal'),
    ('【正式版設定】', 'Heading 3'),
    ('各池需備有足夠牌數（每池至少涵蓋該期間天數 × 3 張）。', 'Normal'),
    ('【切換機制說明】', 'Heading 3'),
    ('換日時（advanceDay）自動比對新天數與閾值，若跨越閾值則將新年齡池的所有牌洗牌後附加至 mainQueue 末端。', 'Normal'),
    ('舊年齡池尚未打完的牌不移除，繼續留在 queue 中直到被抽到。', 'Normal'),
    ('每個年齡池的牌只在切換時加入一次，不重複補充。', 'Normal'),
    ('【目前各池牌數（測試階段）】', 'Heading 3'),
    ('童年：5 張（card_001, 006~009）', 'Normal'),
    ('少年：3 張（card_010~012）', 'Normal'),
    ('青年：2 張（card_002~003）', 'Normal'),
    ('中年：2 張（card_004~005）', 'Normal'),
    ('老年：0 張（待補充）', 'Normal'),
    ('', 'Normal'),
]

ref = three_blank
for text, style in reversed(rows_3):
    ref = insert_after(ref, text, style)
print('三章 完成')

# ── 四、重要常數：找空行後插入 ───────────────────────────────
four_blank = None
for i, p in enumerate(doc.paragraphs):
    if '四、重要常數' in p.text:
        four_blank = doc.paragraphs[i + 1]._element
        break

rows_4 = [
    ('AGE_POOL_THRESHOLDS（年齡池切換閾值）', 'Heading 3'),
    ('定義在 src/core/constants/gameConfig.ts。', 'Normal'),
    ('測試階段：{ childhood:0, juvenile:3, youth:6, midlife:9, elder:12 }', 'Normal'),
    ('正式版本：{ childhood:0, juvenile:10, youth:30, midlife:60, elder:90 }', 'Normal'),
    ('上線前須將閾值改回正式版數值，並確認各池牌數足夠。', 'Normal'),
    ('', 'Normal'),
]

ref2 = four_blank
for text, style in reversed(rows_4):
    ref2 = insert_after(ref2, text, style)
print('四章 完成')

doc.save(path)
print('Done')
