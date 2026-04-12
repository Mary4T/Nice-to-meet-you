# -*- coding: utf-8 -*-
"""
在補充設計決策.docx 的「七、牌卡格式與效果系統」末尾（7-6 之後、八 之前）
新增 7-7 條件選項系統。
"""
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

path = r'D:\APP\Nice to meet you\設計文件\補充設計決策.docx'
doc = Document(path)

# 建立 style map（避免 duplicate name 問題）
style_map = {}
for s in doc.styles:
    if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name not in style_map:
        style_map[s.name] = s

def find_heading_index(doc, keyword):
    for i, p in enumerate(doc.paragraphs):
        if keyword in p.text and p.style and p.style.name.startswith('Heading'):
            return i
    return -1

def insert_para_after(doc, ref_idx, text, style_name=None):
    ref = doc.paragraphs[ref_idx]._element
    style_obj = style_map.get(style_name) if style_name else None
    p = doc.add_paragraph(text, style=style_obj)
    ref.addnext(p._element)
    return p

# 找插入位置：在「八、靈魂成長系統」之前
# 新內容插在 7-6 末尾的空白行後面，也就是「八」標題的正前方
idx_8 = find_heading_index(doc, '八、靈魂成長系統')
if idx_8 == -1:
    # 備用：找 8-0
    idx_8 = find_heading_index(doc, '8-0')
print(f'插入位置（八的前一段）: {idx_8}')

# 倒序插入
new_content = [
    ('7-7 條件選項系統（Conditional Options）', 'Heading 2'),
    ('7-7-1 設計動機', 'Heading 3'),
    ('部分牌卡在特定數值條件下（如裏生命值極高或壓抑熵極高），玩家應面對程度更深、更極端的選擇。條件選項系統允許同一方向的選項在滿足條件時自動替換為條件版本。', None),
    ('7-7-2 Excel 填寫格式', 'Heading 3'),
    ('在「卡牌資料.xlsx」的「條件」欄填入條件字串；空白 = 基本選項，有值 = 條件選項。', None),
    ('條件格式：Key運算子數值，多條件用 & 連接（AND 邏輯）。', None),
    ('範例：Vitality>30　/　Vitality>30&Entropy>=80', None),
    ('7-7-3 支援的條件 Key', 'Heading 3'),
    ('Core 數值：Vitality / Stability / Drive / Logic / Emotion', None),
    ('靈魂毒素：Entropy / Hypocrisy / Ruthlessness', None),
    ('表層數值：Reputation / Money', None),
    ('遊戲進度：day（天數）', None),
    ('支援的運算子：>　<　>=　<=　=', None),
    ('7-7-4 優先順序規則', 'Heading 3'),
    ('同一方向可有多個條件選項列，系統由下往上讀取優先序（Excel 中越靠下的條件越優先）。', None),
    ('取最後一個所有條件皆成立的版本；若無條件版本符合，顯示基本選項。', None),
    ('7-7-5 條件選項的效果欄位', 'Heading 3'),
    ('條件選項的效果欄位（聲望、穩定等）與基本選項共用同一組欄位，只是填入不同數值。', None),
    ('條件字串只決定「何時出現」，效果欄位決定「選擇後發生什麼」，兩者完全獨立。', None),
    ('7-7-6 技術實作', 'Heading 3'),
    ('JSON 結構：conditionals 陣列附加於各方向選項內，每項含 condition（條件陣列）、text、effects。', None),
    ('LifetimeScreen 中的 getEffectiveOption() 在每次顯示牌面與確認選擇時即時評估條件。', None),
    ('夢境牌不支援條件選項（設計上不需要）。', None),
    ('OR 條件（或）目前不支援，待有需求再擴充。', None),
    ('', None),
]

ref_idx = idx_8 - 1  # 插在「八」前一段後面
for text, style_name in reversed(new_content):
    insert_para_after(doc, ref_idx, text, style_name)

doc.save(path)
print('Done — 7-7 條件選項系統已新增')
