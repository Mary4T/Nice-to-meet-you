# -*- coding: utf-8 -*-
"""
1. 刪除錯誤放在七章下的 7-7 條件選項段落
2. 在文件末尾新增「十、條件選項系統」獨立章節
樣式完全對齊現有章節（只用 pStyle，不加多餘 rPr）
"""
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from lxml import etree
import copy

path = r'D:\APP\Nice to meet you\設計文件\補充設計決策.docx'
doc = Document(path)

# ── 建立 style_map ──────────────────────────────────────────
style_map = {}
for s in doc.styles:
    if s.type == WD_STYLE_TYPE.PARAGRAPH and s.name not in style_map:
        style_map[s.name] = s

# ── 1. 找出並刪除 7-7 區段 ──────────────────────────────────
# 從「7-7 條件選項系統」一直到「八、靈魂成長系統」之前（不含八）
body = doc.element.body
paras = list(body.iterchildren())

start_del = None
end_del = None   # exclusive

for i, elem in enumerate(paras):
    if elem.tag.endswith('}p'):
        text = ''.join(r.text or '' for r in elem.iter() if r.tag.endswith('}t'))
        if '7-7' in text and '條件選項' in text and start_del is None:
            start_del = i
        if '八、靈魂成長系統' in text and start_del is not None and end_del is None:
            end_del = i

print(f'刪除範圍: paras[{start_del}:{end_del}]')

if start_del is not None and end_del is not None:
    to_remove = paras[start_del:end_del]
    for elem in to_remove:
        body.remove(elem)
    print(f'已刪除 {len(to_remove)} 個段落')
else:
    print('未找到 7-7 段落，跳過刪除')

# ── 2. 在文件末尾新增十章 ────────────────────────────────────
def add_heading(doc, style_map, text, level):
    """新增標題段落，style 只用 pStyle，不加多餘 rPr"""
    style_name = f'Heading {level}'
    p = doc.add_paragraph(style=style_map.get(style_name))
    p.text = text
    return p

def add_body(doc, style_map, text):
    """新增正文段落"""
    p = doc.add_paragraph(style=style_map.get('Normal'))
    p.text = text
    return p

# 空行（Normal 空白段落）
def add_blank(doc, style_map):
    p = doc.add_paragraph(style=style_map.get('Normal'))
    p.text = ''
    return p

content = [
    (1, '十、條件選項系統（Conditional Options）'),
    (2, '10-0 概述'),
    (0, '部分牌卡在特定數值條件下，玩家應面對程度更深、更極端的選擇。條件選項系統允許同一方向的選項在滿足條件時自動替換為條件版本。'),
    (2, '10-1 設計動機'),
    (0, '隨著遊戲進行，玩家的裏生命值、靈魂毒素等數值可能達到極端狀態。若仍顯示與初始相同的選項，會失去敘事張力。條件選項讓牌卡在極端情況下提供「程度更深」的選擇，強化沉浸感與後果重量。'),
    (2, '10-2 Excel 填寫格式'),
    (0, '在「卡牌資料.xlsx」的「條件」欄填入條件字串。'),
    (0, '空白條件 = 基本選項；有值 = 條件選項。'),
    (0, '條件格式：Key 運算子 數值，多條件用 & 連接（AND 邏輯）。'),
    (0, '範例：Vitality>30　/　Vitality>30&Entropy>=80'),
    (2, '10-3 支援的條件 Key'),
    (0, 'Core 數值：Vitality / Stability / Drive / Logic / Emotion'),
    (0, '靈魂毒素：Entropy / Hypocrisy / Ruthlessness'),
    (0, '表層數值：Reputation / Money'),
    (0, '遊戲進度：day（天數）'),
    (0, '支援的運算子：>　<　>=　<=　='),
    (2, '10-4 優先順序規則'),
    (0, '同一方向可有多個條件選項列，系統由下往上讀取優先序（Excel 中越靠下的條件越優先）。'),
    (0, '取最後一個所有條件皆成立的版本；若無條件版本符合，顯示基本選項。'),
    (2, '10-5 條件選項的效果欄位'),
    (0, '條件選項的效果欄位（聲望、穩定等）與基本選項共用同一組欄位，只是填入不同數值。'),
    (0, '條件字串只決定「何時出現」，效果欄位決定「選擇後發生什麼」，兩者完全獨立。'),
    (2, '10-6 技術實作'),
    (0, 'JSON 結構：conditionals 陣列附加於各方向選項內，每項含 condition（條件陣列）、text、effects。'),
    (0, 'LifetimeScreen 中的 getEffectiveOption() 在每次顯示牌面與確認選擇時即時評估條件。'),
    (0, '夢境牌不支援條件選項（設計上不需要）。'),
    (0, 'OR 條件（或）目前不支援，待有需求再擴充。'),
]

for level, text in content:
    if level == 0:
        add_body(doc, style_map, text)
    else:
        add_heading(doc, style_map, text, level)

doc.save(path)
print('Done — 十、條件選項系統 已新增，7-7 已刪除')
